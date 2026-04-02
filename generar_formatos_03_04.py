
# -*- coding: utf-8 -*-
"""
Generador de Formato 03 y Formato 04 — NATURACOR
Proyecto: Pruebas y Calidad de Software
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ──────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, border_color="000000"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        bd = OxmlElement(f'w:{side}')
        bd.set(qn('w:val'), 'single')
        bd.set(qn('w:sz'), '4')
        bd.set(qn('w:space'), '0')
        bd.set(qn('w:color'), border_color)
        tcBorders.append(bd)
    tcPr.append(tcBorders)

def heading(doc, text, level=1, color="2E7D32"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13 if level == 1 else 11)
    run.font.color.rgb = RGBColor.from_string(color)
    return p

def sub_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string("1B5E20")
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    return p

def body(doc, text, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.italic = italic
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    return p

def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_page_break(doc):
    doc.add_page_break()

def add_title_block(doc, formato_num, formato_titulo):
    """Adds a styled title block for each format."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"FORMATO {formato_num:02d}")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string("FFFFFF")
    # Shade paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '2E7D32')
    pPr.append(shd)
    p.paragraph_format.space_after = Pt(2)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(formato_titulo)
    run2.bold = True
    run2.font.size = Pt(13)
    run2.font.color.rgb = RGBColor.from_string("FFFFFF")
    pPr2 = p2._p.get_or_add_pPr()
    shd2 = OxmlElement('w:shd')
    shd2.set(qn('w:val'), 'clear')
    shd2.set(qn('w:color'), 'auto')
    shd2.set(qn('w:fill'), '388E3C')
    pPr2.append(shd2)
    p2.paragraph_format.space_after = Pt(10)

def add_datos_generales(doc):
    """Tabla de Datos Generales reutilizable."""
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    datos = [
        ("Nombre del proyecto", "NATURACOR — Sistema Web Empresarial para la Gestión Integral de Tiendas Naturistas con Módulo de Análisis Inteligente y Recomendación de Productos basado en el Perfil de Salud del Cliente"),
        ("Integrantes del equipo", "BENDEZU LAGOS JACK JOSHUA\nJULCA LAUREANO DICKMAR WILBER\nREYES CORDERO ITALO EDUARDO"),
        ("Nombre de la empresa o institución", "NATURACOR — Tienda de Productos Naturales"),
        ("Nombre del proceso analizado", "Gestión de Atención al Cliente"),
        ("Docente", "Maglioni Arana Caparachin"),
        ("Fecha de elaboración", "01 de abril de 2026"),
    ]
    for i, (label, value) in enumerate(datos):
        row = table.rows[i]
        cell_label = row.cells[0]
        cell_value = row.cells[1]
        set_cell_bg(cell_label, "C8E6C9")
        cell_label.text = label
        cell_label.paragraphs[0].runs[0].bold = True
        cell_label.paragraphs[0].runs[0].font.size = Pt(9.5)
        cell_value.text = value
        cell_value.paragraphs[0].runs[0].font.size = Pt(9.5)
        set_cell_borders(cell_label)
        set_cell_borders(cell_value)
    doc.add_paragraph()

def style_header_row(row, bg="2E7D32"):
    for cell in row.cells:
        set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string("FFFFFF")
                run.font.size = Pt(9.5)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def add_checkbox(doc, text, checked=True):
    p = doc.add_paragraph()
    mark = "☑" if checked else "☐"
    run = p.add_run(f"  {mark}  {text}")
    run.font.size = Pt(10)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p

# ══════════════════════════════════════════════
# DOCUMENTO PRINCIPAL
# ══════════════════════════════════════════════
doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# Default font
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10)

# ══════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_title.add_run("PROYECTO NATURACOR")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor.from_string("2E7D32")

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p_sub.add_run("Pruebas y Calidad de Software")
run2.bold = True
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor.from_string("555555")

doc.add_paragraph()
p_formats = doc.add_paragraph()
p_formats.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_formats.add_run("FORMATO 03 — Diagrama BPM (Modelado del Proceso)\nFORMATO 04 — Identificación de Problemas del Proceso")
r.font.size = Pt(12)
r.bold = True
r.font.color.rgb = RGBColor.from_string("388E3C")

doc.add_paragraph()
doc.add_paragraph()
p_team = doc.add_paragraph()
p_team.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p_team.add_run(
    "Integrantes:\nBENDEZU LAGOS JACK JOSHUA\n"
    "JULCA LAUREANO DICKMAR WILBER\nREYES CORDERO ITALO EDUARDO\n\n"
    "Docente: Maglioni Arana Caparachin\n"
    "Fecha: 01 de abril de 2026"
)
r2.font.size = Pt(11)

add_page_break(doc)

# ══════════════════════════════════════════════
# FORMATO 03 — DIAGRAMA BPM
# ══════════════════════════════════════════════
add_title_block(doc, 3, "Diagrama BPM (Modelado del Proceso)")

# 1. DATOS GENERALES
heading(doc, "1. Datos Generales del Proyecto")
add_datos_generales(doc)

# 2. DESCRIPCIÓN GENERAL DEL PROCESO
heading(doc, "2. Descripción General del Proceso")
body(doc,
    "El proceso analizado es la Gestión de Atención al Cliente en las sucursales de NATURACOR, "
    "tienda de productos naturales y cordiales en Perú. Este proceso es el núcleo operativo del negocio, "
    "pues define la calidad del servicio que recibe cada cliente desde su llegada hasta la conclusión de "
    "la interacción, ya sea mediante una venta, la resolución de una consulta o el registro de un reclamo.")
body(doc,
    "Propósito: Garantizar una atención personalizada, segura y eficiente al cliente, apoyada en el "
    "sistema NATURACOR (módulo Recetario y Asistente IA), de modo que cada cliente reciba el "
    "asesoramiento correcto en productos naturales según su perfil de salud.")
body(doc,
    "Alcance:\n"
    "  • INICIO: Llegada física del cliente a cualquiera de las dos sucursales de NATURACOR.\n"
    "  • FIN: Cierre de la interacción — venta completada, consulta resuelta sin compra, o reclamo "
    "formalmente registrado para seguimiento.")
body(doc,
    "Contexto: NATURACOR opera con dos sucursales físicas. Cada sucursal cuenta con un vendedor "
    "que gestiona el proceso de atención apoyándose en el sistema web NATURACOR (Laravel 12). "
    "Los reclamos complejos son escalados al Administrador/Propietaria. El sistema registra las ventas, "
    "consulta el Recetario y genera boletas de manera automática.")

# 3. DIAGRAMA BPM AS-IS
heading(doc, "3. Diagrama BPM del Proceso Actual (AS-IS)")
body(doc,
    "El siguiente diagrama representa el flujo actual (AS-IS) del proceso de Gestión de Atención al "
    "Cliente en NATURACOR, elaborado con notación BPMN estándar. El diagrama fue construido con "
    "draw.io y muestra los cuatro Pools/Lanes: Cliente, Vendedor, Sistema NATURACOR y Administrador.")

body(doc,
    "[Ver imagen del Diagrama BPM adjunta — elaborada en draw.io con notación BPMN 2.0]",
    italic=True)

body(doc,
    "Descripción sintetizada del flujo AS-IS:\n"
    "  1. Cliente llega a la sucursal.\n"
    "  2. Vendedor verifica si el cliente está registrado (búsqueda por DNI).\n"
    "  3. Si no está registrado → se ofrece registro en el sistema.\n"
    "  4. Vendedor identifica la dolencia o necesidad del cliente.\n"
    "  5. Sistema NATURACOR es consultado (módulo Recetario).\n"
    "  6. Vendedor asesora y recomienda productos adecuados.\n"
    "  7. Compuerta: ¿El cliente presenta consulta o reclamo?\n"
    "     a. Consulta → Vendedor responde con apoyo del Asistente IA.\n"
    "     b. Reclamo complejo → Escalar al Administrador.\n"
    "  8. Compuerta: ¿El cliente acepta el producto?\n"
    "     a. SÍ → Registrar venta en POS → Seleccionar método de pago → Sistema genera Boleta "
    "B001-XXXXXX → Entregar producto y comprobante.\n"
    "     b. NO → Agradecer y despedir al cliente.\n"
    "  9. FIN del proceso.")

# 4. ELEMENTOS BPMN UTILIZADOS
heading(doc, "4. Elementos BPMN Utilizados")
cols = ["N°", "Elemento BPMN", "Descripción", "Uso en el proceso"]
rows_data = [
    ("1", "Evento Inicio", "Círculo simple de borde fino. Marca el comienzo del proceso.",
     "Arranca cuando el cliente llega físicamente a la sucursal NATURACOR."),
    ("2", "Evento Fin", "Círculo de borde grueso. Marca el término del proceso.",
     "Se utiliza en dos puntos: (a) tras la entrega del producto y (b) tras agradecer y despedir al cliente sin venta."),
    ("3", "Tarea (Task)", "Rectángulo con esquinas redondeadas. Representa una actividad realizada por un actor.",
     "Recepción del cliente, Identificar necesidad, Asesorar y recomendar, Registrar venta en POS, Seleccionar método de pago, Generar Boleta, Entregar producto, Responder con apoyo IA, Escalar al Administrador."),
    ("4", "Compuerta Exclusiva (XOR Gateway)", "Rombo con aspa (X). Solo uno de los caminos se ejecuta.",
     "¿El cliente está registrado en el sistema? / ¿El cliente presenta consulta o reclamo? / ¿El cliente acepta el producto?"),
    ("5", "Flujo de Secuencia", "Flecha sólida. Conecta los elementos en orden.",
     "Conecta todas las actividades y compuertas del proceso de manera ordenada."),
    ("6", "Pool", "Contenedor principal (rectángulo grande). Agrupa todo el proceso.",
     "Pool principal: PROCESO A — Gestión de Atención al Cliente (AS-IS) : NATURACOR"),
    ("7", "Lane", "Subdivisión horizontal dentro del Pool. Identifica al actor responsable.",
     "Lanes: Cliente, Vendedor, Sistema NATURACOR, Administrador."),
    ("8", "Anotación de texto", "Cuadro de texto con corchete. Agrega información adicional.",
     "Etiquetas Sí/No en compuertas y nombre del flujo de reclamo ('Reclamo')."),
]
table4 = doc.add_table(rows=len(rows_data)+1, cols=4)
table4.style = 'Table Grid'
table4.alignment = WD_TABLE_ALIGNMENT.CENTER
# Header
hdr = table4.rows[0]
for i, col in enumerate(cols):
    hdr.cells[i].text = col
style_header_row(hdr)
# Widths
widths = [Cm(1), Cm(3.5), Cm(6), Cm(7.5)]
for i, row in enumerate(table4.rows):
    for j, cell in enumerate(row.cells):
        cell.width = widths[j]
for i, data in enumerate(rows_data):
    row = table4.rows[i+1]
    for j, val in enumerate(data):
        row.cells[j].text = val
        row.cells[j].paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_borders(row.cells[j])
    if i % 2 == 1:
        for cell in row.cells:
            set_cell_bg(cell, "F1F8E9")
doc.add_paragraph()

# 5. IDENTIFICACIÓN DE ACTORES
heading(doc, "5. Identificación de Actores (Pools / Lanes)")
cols5 = ["N°", "Actor", "Tipo\n(Interno / Externo / Sistema)", "Lane asignado"]
rows5 = [
    ("1", "Cliente", "Externo", "Lane: Cliente"),
    ("2", "Vendedor / Empleado", "Interno", "Lane: Vendedor"),
    ("3", "Administrador / Propietaria", "Interno", "Lane: Administrador"),
    ("4", "Sistema NATURACOR", "Sistema", "Lane: Sistema NATURACOR"),
    ("5", "Asistente IA (Gemini API)", "Sistema Externo", "Lane: Sistema NATURACOR\n(subproceso IA dentro del sistema)"),
]
table5 = doc.add_table(rows=len(rows5)+1, cols=4)
table5.style = 'Table Grid'
table5.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr5 = table5.rows[0]
for i, col in enumerate(cols5):
    hdr5.cells[i].text = col
style_header_row(hdr5)
for i, data in enumerate(rows5):
    row = table5.rows[i+1]
    for j, val in enumerate(data):
        row.cells[j].text = val
        row.cells[j].paragraphs[0].runs[0].font.size = Pt(9.5)
        set_cell_borders(row.cells[j])
    if i % 2 == 1:
        for cell in row.cells:
            set_cell_bg(cell, "F1F8E9")
doc.add_paragraph()

# 6. DESCRIPCIÓN DEL FLUJO DEL PROCESO
heading(doc, "6. Descripción del Flujo del Proceso")
steps = [
    ("Paso 1 — Llegada del cliente (Lane: Cliente)",
     "El cliente llega físicamente a una de las sucursales de NATURACOR. Este evento dispara el inicio "
     "del proceso de atención."),
    ("Paso 2 — Recepción y registro (Lane: Vendedor + Sistema NATURACOR)",
     "El vendedor saluda al cliente y verifica en el sistema si está registrado mediante búsqueda por DNI.\n"
     "  • Compuerta XOR 1: ¿El cliente está registrado?\n"
     "    – SÍ: Se recupera su historial de compras y perfil.\n"
     "    – NO: El vendedor ofrece el registro en el sistema NATURACOR para fidelización."),
    ("Paso 3 — Identificación de necesidades (Lane: Vendedor)",
     "El vendedor escucha activamente la dolencia o necesidad del cliente (producto específico, "
     "condición de salud, consulta general)."),
    ("Paso 4 — Consulta al Recetario (Lane: Sistema NATURACOR)",
     "El sistema NATURACOR es consultado en el módulo Recetario, que asocia enfermedades/condiciones "
     "con productos naturales recomendados y restringidos. Este es el proceso de asesoramiento inteligente."),
    ("Paso 5 — Asesoramiento y recomendación (Lane: Vendedor)",
     "El vendedor presenta al cliente los productos adecuados (precio, beneficio, modo de uso), "
     "verificando stock disponible en el sistema.\n"
     "  • Compuerta XOR 2: ¿El cliente presenta consulta o reclamo?\n"
     "    – CONSULTA: El vendedor responde con apoyo del Asistente IA (Gemini API) cuando hay internet.\n"
     "    – RECLAMO COMPLEJO: Se escala al Administrador (Lane: Administrador) para resolución.\n"
     "    – SIN INCIDENCIA: El flujo continúa directamente a la decisión de compra."),
    ("Paso 6 — Decisión de compra (Lane: Cliente)",
     "  • Compuerta XOR 3: ¿El cliente acepta el producto?\n"
     "    – NO: El vendedor agradece y despide al cliente. FIN 1.\n"
     "    – SÍ: Continúa al registro de la venta."),
    ("Paso 7 — Registro de la venta (Lane: Vendedor + Sistema NATURACOR)",
     "El vendedor registra los productos en el POS del sistema. El sistema calcula subtotal, "
     "IGV (18% incluido en el precio peruano) y total."),
    ("Paso 8 — Cobro y emisión de boleta (Lane: Vendedor + Sistema NATURACOR)",
     "El cliente selecciona el método de pago: Efectivo, Yape o Plin. El sistema genera "
     "automáticamente la boleta correlativa B001-XXXXXX como comprobante de pago."),
    ("Paso 9 — Entrega y cierre (Lane: Vendedor → Cliente)",
     "El vendedor entrega el producto junto con la boleta. Refuerza las instrucciones de uso y despide "
     "cordialmente al cliente. FIN 2."),
]
for title, desc in steps:
    sub_heading(doc, title)
    body(doc, desc)

# 7. VALIDACIÓN DEL MODELO
heading(doc, "7. Validación del Modelo")
body(doc, "Marque con un aspa (X) si se cumple los enunciados:")
checks = [
    "El proceso tiene evento de inicio y fin claramente definidos.",
    "Todas las actividades están conectadas correctamente.",
    "Se utilizan correctamente los elementos BPMN.",
    "Cada actividad tiene un actor asignado.",
    "El flujo es coherente y entendible.",
]
for check in checks:
    add_checkbox(doc, check, checked=True)
body(doc,
    "\nObservación: Todos los ítems de validación son cumplidos por el modelo BPM elaborado. "
    "El diagrama fue revisado por el equipo y validado garantizando coherencia, completitud de actores "
    "y uso correcto de la notación BPMN 2.0.")

# 8. EVIDENCIAS
heading(doc, "8. Evidencias")
body(doc,
    "Se adjunta a continuación la captura del diagrama BPM elaborado en draw.io con notación BPMN 2.0, "
    "mostrando los cuatro Lanes (Cliente, Vendedor, Sistema NATURACOR, Administrador) y el flujo "
    "completo del proceso de Gestión de Atención al Cliente AS-IS:")
body(doc,
    "Captura 1: Diagrama BPM — Gestión de Atención al Cliente (AS-IS) — NATURACOR\n"
    "Herramienta: draw.io (diagrams.net)\n"
    "Notación: BPMN 2.0\n"
    "[Imagen del diagrama BPM adjuntada según indicación del docente]",
    italic=True)

add_page_break(doc)

# ══════════════════════════════════════════════
# FORMATO 04 — IDENTIFICACIÓN DE PROBLEMAS
# ══════════════════════════════════════════════
add_title_block(doc, 4, "Identificación de Problemas del Proceso")

# 1. DATOS GENERALES
heading(doc, "1. Datos Generales del Proyecto")
add_datos_generales(doc)

# 2. DESCRIPCIÓN GENERAL
heading(doc, "2. Descripción General del Proceso")
body(doc,
    "El proceso analizado es la Gestión de Atención al Cliente (AS-IS) de NATURACOR, tienda de "
    "productos naturales y cordiales con dos sucursales físicas en Perú.")
body(doc,
    "Propósito: Brindar atención personalizada al cliente, asesorando en la selección de productos "
    "naturales según su perfil de salud, procesar la venta y emitir el comprobante correspondiente.")
body(doc,
    "Alcance: INICIO: Llegada del cliente a la sucursal. FIN: Entrega del producto y boleta, "
    "resolución de consulta o registro del reclamo.")
body(doc,
    "Contexto organizacional: NATURACOR opera sin módulo de reclamos digitalizado y sin "
    "sincronización entre sus dos sucursales. El proceso depende en gran medida de la competencia "
    "del vendedor y de la disponibilidad de internet para el Asistente IA. No existe mecanismo formal "
    "de medición de satisfacción del cliente (requisito ISO 9001). La tienda ha iniciado su digitalización "
    "con el sistema NATURACOR (Laravel 12), pero persisten cuellos de botella y riesgos operativos "
    "en el proceso AS-IS.")

# 3. LISTADO DE PROBLEMAS
heading(doc, "3. Listado de Problemas Identificados")
cols3 = ["N°", "Actividad del proceso", "Problema identificado", "Tipo de\nproblema",
         "Descripción del problema", "Impacto", "Prioridad\n(Alta/Media/Baja)"]
rows3 = [
    ("1",
     "Asesoramiento y recomendación de productos",
     "Dependencia del conocimiento del vendedor",
     "Calidad",
     "La calidad del asesoramiento depende del nivel de experiencia del vendedor. El módulo Recetario mitiga el riesgo, pero no garantiza una recomendación completa si el vendedor no lo usa correctamente.",
     "Recomendaciones incorrectas, riesgo para la salud del cliente",
     "Alta"),
    ("2",
     "Gestión de reclamos",
     "Reclamos no digitalizados ni con trazabilidad",
     "Control",
     "No existe módulo de reclamos en el sistema. Los incidentes se gestionan verbalmente. No hay historial, seguimiento ni análisis de tendencias para mejora continua.",
     "Pérdida de información, incumplimiento ISO 9001, reincidencia de problemas",
     "Alta"),
    ("3",
     "Identificación del cliente",
     "Sin sincronización entre sucursales",
     "Control / Redundancia",
     "Un cliente registrado en la Sucursal A no puede ser consultado desde la Sucursal B. Cada sede opera de manera independiente, lo que obliga a duplicar el registro del cliente.",
     "Doble registro, mala experiencia del cliente, pérdida del historial de compras",
     "Alta"),
    ("4",
     "Asesoramiento IA (Asistente Gemini)",
     "Dependencia de conectividad a internet",
     "Tiempo / Calidad",
     "El Asistente IA opera en modo offline sin internet, limitando la calidad y profundidad de las recomendaciones avanzadas. El modo offline no tiene la misma capacidad de análisis.",
     "Asesoramiento reducido o incompleto en caso de fallas de internet",
     "Media"),
    ("5",
     "Cierre de la atención al cliente",
     "Ausencia de encuesta de satisfacción del cliente",
     "Control",
     "No existe ningún mecanismo (digital o físico) para medir la satisfacción del cliente tras la atención. Esto incumple el requisito de ISO 9001 de retroalimentación con el cliente.",
     "Sin datos para mejora continua, incumplimiento ISO 9001",
     "Media"),
    ("6",
     "Registro de la venta (POS)",
     "Sin descuentos por fidelización en el POS",
     "Calidad / Tiempo",
     "El POS registra ventas, pero no aplica automáticamente descuentos por fidelización del cliente (basados en historial de compras). El cálculo es manual si aplica.",
     "Demoras en el cobro, errores en descuentos, menor fidelización del cliente",
     "Media"),
    ("7",
     "Identificación del cliente",
     "Registro manual de clientes nuevos genera demoras",
     "Tiempo",
     "Cuando el cliente no está registrado, el vendedor debe ingresar los datos manualmente en el sistema durante la atención, generando colas y demoras innecesarias.",
     "Tiempo de espera elevado, mala experiencia del cliente",
     "Baja"),
    ("8",
     "Escalado de reclamos al Administrador",
     "Sin definición de criterios de escalado",
     "Control",
     "No existe un criterio formal para determinar cuándo un reclamo debe escalarse al Administrador. La decisión es subjetiva y depende del vendedor.",
     "Reclamos complejos no resueltos a tiempo o escalados innecesariamente",
     "Baja"),
]
table3 = doc.add_table(rows=len(rows3)+1, cols=7)
table3.style = 'Table Grid'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr3 = table3.rows[0]
for i, col in enumerate(cols3):
    hdr3.cells[i].text = col
style_header_row(hdr3)
for i, data in enumerate(rows3):
    row = table3.rows[i+1]
    for j, val in enumerate(data):
        row.cells[j].text = val
        row.cells[j].paragraphs[0].runs[0].font.size = Pt(8.5)
        set_cell_borders(row.cells[j])
    if i % 2 == 1:
        for cell in row.cells:
            set_cell_bg(cell, "F1F8E9")
    # Color prioridad
    prioridad_cell = row.cells[6]
    prioridad = data[6].strip().lower()
    if prioridad == "alta":
        set_cell_bg(prioridad_cell, "FFCDD2")
    elif prioridad == "media":
        set_cell_bg(prioridad_cell, "FFF9C4")
    elif prioridad == "baja":
        set_cell_bg(prioridad_cell, "C8E6C9")
doc.add_paragraph()

# 4. CLASIFICACIÓN DE PROBLEMAS
heading(doc, "4. Clasificación de Problemas")

sub_heading(doc, "• Problemas de tiempo (retrasos, cuellos de botella)")
body(doc,
    "  – Problema 7: El registro manual de clientes nuevos durante la atención genera cuellos de "
    "botella y colas en la tienda, especialmente en horas pico.\n"
    "  – Problema 4: La degradación del Asistente IA en modo offline ralentiza el asesoramiento ya "
    "que el vendedor debe recurrir únicamente al Recetario sin sugerencias dinámicas.")

sub_heading(doc, "• Problemas de calidad (errores, reprocesos)")
body(doc,
    "  – Problema 1: La dependencia del conocimiento del vendedor introduce variabilidad en la calidad "
    "del asesoramiento, con riesgo de recomendar productos inadecuados para la condición de salud.\n"
    "  – Problema 6: La falta de aplicación automática de descuentos de fidelización puede generar "
    "errores en el cobro final.")

sub_heading(doc, "• Problemas de control (falta de supervisión o validación)")
body(doc,
    "  – Problema 2: Los reclamos no son digitalizados, lo que imposibilita el seguimiento, el análisis "
    "de causas raíz y el cumplimiento de ISO 9001.\n"
    "  – Problema 5: La ausencia de encuesta de satisfacción impide medir el desempeño real del "
    "proceso y contrastar con los objetivos de calidad.\n"
    "  – Problema 8: Sin criterios formales de escalado, la gestión de reclamos complejos es arbitraria "
    "y puede comprometer la relación con el cliente.")

sub_heading(doc, "• Problemas de redundancia (actividades duplicadas)")
body(doc,
    "  – Problema 3: La falta de sincronización entre sucursales obliga al registro duplicado de un mismo "
    "cliente en ambas sedes, generando datos inconsistentes y esfuerzo redundante.")

sub_heading(doc, "• Otros")
body(doc,
    "  – Riesgo para la salud del cliente (Problema 1): Una recomendación incorrecta de producto "
    "natural puede agravar condiciones médicas preexistentes. Este es el riesgo de mayor impacto "
    "en el negocio y requiere prioridad en la validación del módulo Recetario y IA.\n"
    "  – Incumplimiento normativo ISO 9001 (Problemas 2 y 5): La falta de trazabilidad de reclamos y "
    "medición de satisfacción impide alinear el proceso a estándares internacionales de gestión de calidad.")

# 5. ANÁLISIS DE CAUSAS
heading(doc, "5. Análisis de Causas")
cols5c = ["N°", "Problema identificado", "Causa principal", "Descripción de la causa"]
rows5c = [
    ("1",
     "Dependencia del conocimiento del vendedor",
     "Ausencia de validación obligatoria del Recetario antes de recomendar",
     "El sistema no obliga al vendedor a consultar el módulo Recetario antes de recomendar un producto. El uso es voluntario, lo que deja la calidad del asesoramiento sujeta al criterio individual del empleado."),
    ("2",
     "Reclamos no digitalizados",
     "Falta de módulo de gestión de reclamos en NATURACOR",
     "El sistema NATURACOR no cuenta con un módulo dedicado al registro, seguimiento y resolución de reclamos y quejas del cliente. No fue considerado en la fase de desarrollo inicial."),
    ("3",
     "Sin sincronización entre sucursales",
     "Arquitectura de base de datos sin replicación entre sedes",
     "Cada sucursal opera con su propia instancia de datos. No se ha implementado sincronización en tiempo real ni centralización de la base de datos para ambas sedes."),
    ("4",
     "Dependencia de conectividad para IA",
     "La API Gemini requiere conexión a internet estable",
     "El Asistente IA depende de la API externa de Google Gemini, que requiere internet activo. El modo offline utiliza lógica local limitada, sin la capacidad de análisis contextual avanzado de la IA en la nube."),
    ("5",
     "Ausencia de encuesta de satisfacción",
     "Falta de módulo de retroalimentación del cliente en el sistema",
     "No existe ningún componente en NATURACOR que permita capturar la percepción del cliente tras la atención. No fue contemplado como requerimiento durante el análisis inicial del sistema."),
    ("6",
     "Sin descuentos automáticos por fidelización",
     "Módulo de fidelización no integrado al POS",
     "El sistema registra el historial de compras del cliente, pero no tiene reglas de negocio configuradas para calcular y aplicar descuentos automáticos según el nivel de fidelización en el momento del cobro."),
    ("7",
     "Registro manual de clientes nuevos genera demoras",
     "Proceso de registro no simplificado ni anticipatorio",
     "No existe un mecanismo de pre-registro (QR, ficha digital, WhatsApp) que permita al cliente ingresar sus datos antes de llegar al mostrador, lo que obliga al vendedor a registrarlo manualmente durante la atención."),
    ("8",
     "Sin criterios formales de escalado de reclamos",
     "Ausencia de política y protocolo de atención de reclamos",
     "La tienda no ha definido un protocolo formal (documentado o en el sistema) que establezca con claridad qué tipo de reclamos pueden resolverse por el vendedor y cuáles deben ser escalados al Administrador."),
]
table5c = doc.add_table(rows=len(rows5c)+1, cols=4)
table5c.style = 'Table Grid'
table5c.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr5c = table5c.rows[0]
for i, col in enumerate(cols5c):
    hdr5c.cells[i].text = col
style_header_row(hdr5c)
widths5c = [Cm(1), Cm(4), Cm(4.5), Cm(8.5)]
for i, row in enumerate(table5c.rows):
    for j, cell in enumerate(row.cells):
        cell.width = widths5c[j]
for i, data in enumerate(rows5c):
    row = table5c.rows[i+1]
    for j, val in enumerate(data):
        row.cells[j].text = val
        row.cells[j].paragraphs[0].runs[0].font.size = Pt(8.5)
        set_cell_borders(row.cells[j])
    if i % 2 == 1:
        for cell in row.cells:
            set_cell_bg(cell, "F1F8E9")
doc.add_paragraph()

# 6. RELACIÓN CON EL DIAGRAMA BPM
heading(doc, "6. Relación con el Diagrama BPM")
rel_data = [
    ("Problema 1 — Dependencia del vendedor",
     "Tarea 'Asesorar y recomendar producto(s) adecuado(s)' — Lane: Vendedor. "
     "El diagrama AS-IS no muestra ningún paso de validación obligatoria por parte del sistema antes "
     "de presentar la recomendación al cliente."),
    ("Problema 2 — Reclamos no digitalizados",
     "Compuerta XOR 2 — Rama 'Reclamo → Escalar al administrador' — Lane: Administrador. "
     "El flujo termina con el escalado pero no hay ninguna tarea posterior de 'Registrar reclamo en el sistema'."),
    ("Problema 3 — Sin sincronización entre sucursales",
     "Tarea 'Buscar cliente por DNI en NATURACOR' — Lane: Sistema NATURACOR. "
     "La tarea opera sobre la base de datos local de la sucursal activa, sin acceso a la base de datos "
     "de la otra sede."),
    ("Problema 4 — Dependencia de IA con internet",
     "Tarea 'Responder con apoyo del sistema / Asistente IA' — Lane: Sistema NATURACOR. "
     "El diagrama no diferencia entre el flujo online (IA activa) y offline (modo degradado), lo que "
     "hace invisible este riesgo operativo."),
    ("Problema 5 — Sin encuesta de satisfacción",
     "Evento Fin — Tras la entrega del producto. No existe ninguna tarea previa al cierre que capte "
     "la opinión del cliente sobre la atención recibida."),
    ("Problema 6 — Sin descuentos automáticos",
     "Tarea 'Registrar venta en POS de NATURACOR' — Lane: Vendedor. "
     "El diagrama no refleja una tarea de verificación de descuentos por fidelización antes de pasar al cobro."),
    ("Problema 7 — Registro manual genera demoras",
     "Compuerta XOR 1 — Rama 'NO → Ofrecer registro en el sistema' — Lane: Vendedor. "
     "El flujo solo prevé una bifurcación simple, sin mostrar el tiempo real que implica el registro "
     "manual de datos durante la atención."),
    ("Problema 8 — Sin criterios de escalado",
     "Compuerta XOR 2 — Rama 'Reclamo → Escalar al administrador' — Lane: Administrador. "
     "El criterio de escalado no está definido en el diagrama ni en el sistema; la decisión queda "
     "en manos del vendedor sin soporte formal."),
]
for title, desc in rel_data:
    sub_heading(doc, title)
    body(doc, desc)

# 7. CONCLUSIONES
heading(doc, "7. Conclusiones del Análisis")
body(doc,
    "El análisis del proceso AS-IS de Gestión de Atención al Cliente en NATURACOR ha revelado ocho "
    "problemas significativos, de los cuales tres son de prioridad ALTA y requieren atención inmediata "
    "en la fase de mejora del proceso (TO-BE):")
body(doc,
    "  1. La dependencia del conocimiento del vendedor para el asesoramiento de productos naturales "
    "es el problema de mayor riesgo, ya que puede afectar directamente la salud del cliente. La solución "
    "pasa por hacer obligatorio el uso del módulo Recetario y fortalecer el rol del Asistente IA.")
body(doc,
    "  2. La falta de un módulo de reclamos digitalizado impide la trazabilidad, el seguimiento y la mejora "
    "continua. Este es un requisito directo de ISO 9001 y su ausencia limita la madurez del proceso.")
body(doc,
    "  3. La ausencia de sincronización entre las dos sucursales de NATURACOR genera duplicación de "
    "datos, inconsistencias en el historial del cliente e ineficiencias operativas.")
body(doc,
    "Los demás problemas (prioridad MEDIA y BAJA) impactan en la calidad del servicio, la eficiencia "
    "del proceso y el cumplimiento normativo, y deben ser atendidos en las siguientes fases del proyecto.")
body(doc,
    "En conjunto, estos problemas justifican la necesidad de un proceso TO-BE rediseñado que integre: "
    "(a) validación obligatoria del Recetario, (b) módulo de reclamos en el sistema, (c) sincronización "
    "multisucursal, (d) mecanismo de encuesta de satisfacción, y (e) protocolos formales de escalado "
    "y descuentos por fidelización.")

# 8. EVIDENCIAS
heading(doc, "8. Evidencias")
body(doc,
    "Se adjuntan a continuación las capturas del diagrama BPM donde se evidencian "
    "los problemas identificados:")
evidencias = [
    "Captura 1: Diagrama BPM AS-IS completo — Gestión de Atención al Cliente — NATURACOR (draw.io, BPMN 2.0). Muestra el flujo general con los cuatro Lanes.",
    "Captura 2: Detalle de la Compuerta XOR 2 — Rama de Reclamo → Escalar al Administrador (Problemas 2 y 8): Ausencia de tarea de registro digital del reclamo.",
    "Captura 3: Detalle de la Tarea 'Asesorar y recomendar producto(s)' — Lane Vendedor (Problema 1): No hay tarea obligatoria de consulta al Recetario antes de la recomendación.",
    "Captura 4: Detalle de la Tarea 'Buscar cliente por DNI en NATURACOR' — Lane Sistema NATURACOR (Problema 3): La búsqueda es solo en la base de datos local de la sucursal activa.",
    "Captura 5: Detalle del Evento Fin — Tras entrega del producto (Problema 5): No existe tarea de encuesta de satisfacción antes del cierre del proceso.",
]
for ev in evidencias:
    bullet(doc, ev)

body(doc,
    "\n[Las capturas del diagrama BPM elaborado en draw.io deben adjuntarse según indicación del docente.]",
    italic=True)

# ══════════════════════════════════════════════
# GUARDAR
# ══════════════════════════════════════════════
output_path = r"D:\ESCRITORIO\UNIVERSIDAD\7mo ciclo\PRUEBAS Y CALIDAD DE SOFTWARE\PROYECTO NATURACOR\NATURACOR_Formato03_Formato04.docx"
doc.save(output_path)
print(f"[OK] Documento generado exitosamente en:\n{output_path}")
